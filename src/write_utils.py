"""Shared utilities for writing plain .docx files from paragraph data."""

from __future__ import annotations

import logging

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

STYLE_MAP: dict[str, str] = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Normal": "Normal",
}


def write_plain_docx(paragraphs: list[dict], path: str) -> None:
    """Write a plain .docx without track changes (fallback for no-change case).

    Preserves style, indentation, bold, italic, underline, and font metadata.
    """
    logger.debug("Writing plain .docx to %s (%d paragraphs)", path, len(paragraphs))
    doc = Document()

    for entry in paragraphs:
        p = doc.add_paragraph()

        style_name = entry.get("style_name", "Normal")
        pn = STYLE_MAP.get(style_name, "Normal")
        if hasattr(p, "style"):
            try:
                p.style = doc.styles[pn]
            except KeyError:
                pass

        indent = entry.get("indent_level", 0) or 0
        if indent > 0:
            p.paragraph_format.left_indent = Pt(indent)

        if entry.get("space_before") is not None:
            p.paragraph_format.space_before = Pt(entry["space_before"])
        if entry.get("space_after") is not None:
            p.paragraph_format.space_after = Pt(entry["space_after"])
        if entry.get("line_spacing") is not None:
            p.paragraph_format.line_spacing = entry["line_spacing"]

        run = p.add_run(entry["text"])

        if entry.get("is_bold"):
            run.bold = True
        if entry.get("is_italic"):
            run.italic = True
        if entry.get("is_underline"):
            run.underline = True
        if entry.get("font_size"):
            run.font.size = Pt(entry["font_size"])
        if entry.get("font_name"):
            run.font.name = entry["font_name"]

    doc.save(path)
