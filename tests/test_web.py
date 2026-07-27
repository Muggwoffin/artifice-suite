"""Tests for the FastAPI web frontend (src/web/).

Mocks the LLM call the same way tests/test_llm_client.py does — patching
`src.llm_client._send_request_with_retry` — so no real Ollama/OpenAI/
Anthropic request ever leaves the test process.
"""

from __future__ import annotations

import json
import time
from unittest.mock import patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

import src.web.runtime as runtime
from src.web.server import app


@pytest.fixture(autouse=True)
def isolated_settings(tmp_path, monkeypatch):
    """Never let a test touch the real ~/.artifice_draft/web_settings.json.

    OCR Pipeline (the sibling project this web build is modeled on) shipped a
    test that skipped this once and wrote real content into a developer's
    actual settings file; the fix there was making this the standard fixture
    pattern for every test that touches persisted settings, and this project
    follows the same rule from the start.
    """
    monkeypatch.setattr(runtime, "_SETTINGS_PATH", tmp_path / "web_settings.json")


@pytest.fixture()
def client():
    return TestClient(app)


@pytest.fixture()
def docx_bytes(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello wrold")
    doc.add_paragraph("This is a secnd paragraph.")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p.read_bytes()


def _wait_for_thread(doc_id: str, timeout: float = 5.0) -> None:
    doc = runtime.state.get(doc_id)
    assert doc is not None
    assert doc.thread is not None
    doc.thread.join(timeout=timeout)


def _mock_edits_response():
    return json.dumps([
        {"paragraph_index": 0, "edited_text": "Hello world", "status": "edited"},
        {"paragraph_index": 1, "edited_text": None, "status": "unchanged"},
    ])


# --------------------------------------------------------------------- upload

def test_upload_rejects_non_docx(client):
    res = client.post("/api/upload", files={"file": ("notes.txt", b"hello", "text/plain")})
    assert res.status_code == 400


def test_upload_parses_paragraphs(client, docx_bytes):
    res = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    assert res.status_code == 200
    body = res.json()
    assert body["paragraph_count"] == 2
    assert body["stage"] == "uploaded"
    assert body["doc_id"]


# ------------------------------------------------------------------- settings

def test_settings_roundtrip_and_merge(client):
    res = client.get("/api/settings")
    assert res.status_code == 200
    defaults = res.json()
    assert defaults["llm_provider"] == "ollama"

    res = client.post("/api/settings", json={"batch_size": 3, "enable_review": True})
    assert res.status_code == 200
    updated = res.json()
    assert updated["batch_size"] == 3
    assert updated["enable_review"] is True
    # author_name untouched by the partial patch — merge, not replace.
    assert updated["author_name"] == defaults["author_name"]

    res = client.post("/api/settings", json={"author_name": "Test Author"})
    updated2 = res.json()
    assert updated2["author_name"] == "Test Author"
    # earlier fields survive the second, unrelated patch.
    assert updated2["batch_size"] == 3
    assert updated2["enable_review"] is True


def test_settings_never_expose_api_keys(client):
    res = client.get("/api/settings")
    body = res.json()
    assert "openai_api_key" not in body
    assert "anthropic_api_key" not in body


# ------------------------------------------------------------- run: no review

def test_run_without_review_writes_output_and_is_downloadable(client, docx_bytes):
    client.post("/api/settings", json={"enable_review": False})

    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]

    with patch("src.llm_client._send_request_with_retry", return_value=_mock_edits_response()):
        start = client.post(f"/api/run/{doc_id}/start")
        assert start.status_code == 200
        _wait_for_thread(doc_id)

    status = client.get(f"/api/run/{doc_id}/status").json()
    assert status["stage"] == "done"
    assert status["output_filename"]

    download = client.get(f"/api/run/{doc_id}/download")
    assert download.status_code == 200
    assert len(download.content) > 0


def test_download_404s_before_finalize(client, docx_bytes):
    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]
    res = client.get(f"/api/run/{doc_id}/download")
    assert res.status_code == 404


def test_unknown_doc_id_404s(client):
    assert client.get("/api/run/doesnotexist/status").status_code == 404
    assert client.post("/api/run/doesnotexist/start").status_code == 404
    assert client.get("/api/run/doesnotexist/review").status_code == 404
    assert client.get("/api/run/doesnotexist/download").status_code == 404


# ---------------------------------------------------------------- run: review

def test_review_flow_reject_keeps_original_text(client, docx_bytes):
    client.post("/api/settings", json={"enable_review": True})

    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]

    with patch("src.llm_client._send_request_with_retry", return_value=_mock_edits_response()):
        client.post(f"/api/run/{doc_id}/start")
        _wait_for_thread(doc_id)

    status = client.get(f"/api/run/{doc_id}/status").json()
    assert status["stage"] == "awaiting_review"

    review = client.get(f"/api/run/{doc_id}/review").json()
    # Only the actually-changed paragraph appears — the "unchanged" one from
    # the mocked response is filtered out, same rule cli_review() applies.
    assert len(review["items"]) == 1
    item = review["items"][0]
    assert item["paragraph_index"] == 0
    assert item["original_text"] == "Hello wrold"
    assert item["edited_text"] == "Hello world"
    assert item["diff"]["original_ranges"]

    submit = client.post(f"/api/run/{doc_id}/review", json={
        "decisions": [{"paragraph_index": 0, "approved": False, "replacement_text": None}],
    })
    assert submit.status_code == 200
    assert submit.json()["stage"] == "done"

    download = client.get(f"/api/run/{doc_id}/download")
    assert download.status_code == 200


def test_review_endpoint_409s_when_not_awaiting_review(client, docx_bytes):
    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]
    res = client.get(f"/api/run/{doc_id}/review")
    assert res.status_code == 409
