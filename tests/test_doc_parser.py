"""Tests for doc_parser module."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.doc_parser import parse_docx, write_docx


def _make_test_docx(path: str, paragraphs: list[dict]) -> None:
    """Create a test .docx from paragraph dicts."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p["text"])
    doc.save(path)


def test_parse_simple_doc(tmp_path: Path):
    tmp = tmp_path / "input.docx"
    _make_test_docx(str(tmp), [
        {"text": "Hello world", "style_name": "Normal"},
        {"text": "This is a second paragraph.", "style_name": "Normal"},
    ])

    result = parse_docx(str(tmp))
    assert len(result) == 2
    assert result[0]["text"] == "Hello world"
    assert result[1]["text"] == "This is a second paragraph."


def test_parse_empty_paragraphs_skipped(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("Keep this")
    p = tmp_path / "empty.docx"
    doc.save(str(p))

    result = parse_docx(str(p))
    assert len(result) == 1
    assert result[0]["text"] == "Keep this"


def test_write_roundtrip(tmp_path: Path):
    doc = Document()
    p1 = doc.add_paragraph("First paragraph")
    p1.style.name = "Normal"
    p2 = doc.add_paragraph("Second paragraph")
    p2.paragraph_format.left_indent = Pt(36)

    tmp_in = tmp_path / "roundtrip_in.docx"
    tmp_out = tmp_path / "roundtrip_out.docx"
    doc.save(str(tmp_in))

    parsed = parse_docx(str(tmp_in))
    assert len(parsed) >= 2

    write_docx(parsed, str(tmp_out))
    assert tmp_out.exists()


def test_parse_nonexistent_raises():
    try:
        parse_docx("does_not_exist.docx")
    except ValueError as e:
        assert "Could not open" in str(e)


def test_parse_bold_italic_detection(tmp_path: Path):
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Bold and italic text")
    run.bold = True
    run.italic = True
    doc.add_paragraph("Plain text")

    path = tmp_path / "formatting.docx"
    doc.save(str(path))

    result = parse_docx(str(path))
    assert len(result) == 2
    assert result[0]["is_bold"] is True
    assert result[0]["is_italic"] is True
    assert result[1]["is_bold"] is False
    assert result[1]["is_italic"] is False


def test_parse_style_name(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Normal text")
    heading = doc.add_paragraph("Heading text")
    heading.style = doc.styles["Heading 1"]

    path = tmp_path / "styles.docx"
    doc.save(str(path))

    result = parse_docx(str(path))
    assert len(result) == 2
    assert result[0]["style_name"] == "Normal"
    assert result[1]["style_name"] == "Heading 1"
