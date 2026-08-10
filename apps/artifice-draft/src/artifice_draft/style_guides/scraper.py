# SPDX-FileCopyrightText: 2026 Maurice Casey
#
# SPDX-License-Identifier: AGPL-3.0-or-later

"""Scrape a journal style guide from a URL and convert it to a StyleGuide.

Uses readability-lxml to extract the main article content from any HTML page,
then asks the LLM to parse the extracted text into the StyleGuide schema.
The result is returned (not saved) so the caller can show a review step.
"""

from __future__ import annotations

import asyncio
import ipaddress
import logging
import re
import socket
from urllib.parse import urljoin, urlparse

import requests
from pydantic import BaseModel, Field
from requests.adapters import HTTPAdapter

from artifice_draft.config import AppConfig
from artifice_draft.models import LLMProvider
from artifice_draft.style_guides.base import StyleGuide

from model_harness.anthropic_adapter import AnthropicProvider
from model_harness.contract import (
    HarnessResult,
    ModelConnectorConfig,
    SchemaValidationFailed,
    StructuredOutputMode,
    StructuredOutputUnsupported,
    StructuredRequest,
)
from model_harness.driver import run_structured
from model_harness.endpoint_policy import EndpointPolicy as ConcreteEndpointPolicy
from model_harness.openai_adapter import OpenAIProvider

logger = logging.getLogger(__name__)

# --- rate-limit / safety constants ----------------------------------------- #
_MAX_RESPONSE_BYTES = 5 * 1024 * 1024  # 5 MB — journal guidelines are small
_REQUEST_TIMEOUT = 30  # seconds
_MAX_REDIRECTS = 5  # prevent redirect loops
_USER_AGENT = "ArtificeDraft/1.0 (+https://github.com/anomalyco/opencode) style-guide-scraper"

# --- Wire schema for harness extraction ------------------------------------ #


class _GuideExtractionShape(BaseModel):
    """Pydantic model mirroring :class:`StyleGuide` for harness validation.

    Every field has a default so the harness can validate partial responses.
    """

    name: str = ""
    edition: str = ""
    citation_style: str = ""
    footnote_format: str = ""
    bibliography_format: str = ""
    heading_capitalization: str = ""
    prose_rules: list[str] = Field(default_factory=list)
    quotation_rules: str = ""
    abbreviation_rules: str = ""
    date_format: str = ""
    page_reference_format: str = ""
    url_format: str = ""
    system_prompt_addendum: str = ""
    custom_rules: list[str] = Field(default_factory=list)


# The LLM prompt that asks it to parse extracted text into StyleGuide JSON.
_EXTRACTION_SYSTEM_PROMPT = """\
You are a text analyst. You will receive the main body text of a journal's \
author guidelines or style guide webpage. Your task is to extract the editing \
rules and formatting conventions and return a JSON object matching this schema:

{
  "name": "string — the journal or style guide name",
  "edition": "string — edition or version, if stated",
  "citation_style": "string — e.g. 'notes-bibliography', 'author-date', 'MLA'",
  "footnote_format": "string — rules for footnote formatting, with examples if available",
  "bibliography_format": "string — rules for bibliography/reference list formatting",
  "heading_capitalization": "string — e.g. 'title-case', 'sentence-case'",
  "prose_rules": ["string — individual prose/style rules, one per element"],
  "quotation_rules": "string — rules for direct quotations",
  "abbreviation_rules": "string — rules for abbreviations and acronyms",
  "date_format": "string — how dates should be formatted",
  "page_reference_format": "string — how page numbers/references should appear",
  "url_format": "string — rules for URLs and DOIs",
  "system_prompt_addendum": "string — a concise summary of ALL the key editing rules, \
written as clear instructions an LLM editor should follow. This is the most important \
field: it should distill the entire style guide into actionable rules.",
  "custom_rules": ["string — any rules that don't fit the above categories"]
}

Rules:
- Return ONLY valid JSON, no markdown fences, no commentary.
- If a field is not addressed in the source text, use an empty string (or empty array).
- The system_prompt_addendum should be thorough — it is the primary text the editor \
LLM will read. Include as many concrete rules as the source text supports.
- For prose_rules, extract each distinct rule as a separate list element.
- If the source text is not a style guide at all, return an object with only \
"name" set to a description of what the page contained, and all other fields empty.
"""


# ---------------------------------------------------------------------------
# URL validation for user-supplied style-guide URLs
# ---------------------------------------------------------------------------


def _validate_public_url(url: str) -> str:
    """Validate that *url* points to a public (non-internal) host.

    This is the opposite of ``EndpointPolicy`` — it denies loopback,
    private-network, and link-local addresses, and permits public ones.
    ``EndpointPolicy`` governs model-endpoint connections; this governs
    user-supplied web pages, where a loopback address would be an SSRF
    rather than a legitimate local model.

    Returns *url* unchanged on success; raises ``ValueError`` otherwise.
    """
    _resolve_public_host(url)
    return url


def _resolve_public_host(url: str) -> tuple[str, str]:
    """Validate *url* as public and return (hostname, pinned_ip_address).

    Performs DNS resolution, rejects non-global addresses, and returns one
    validated public IP so the caller can pin the follow-on TCP connection.
    Separating resolution from the connection closes the TOCTOU race where
    a short-TTL DNS record presents a public address at validation time and
    a private one at connect time.
    """
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        raise ValueError(f"URL must use http or https, got: {parsed.scheme or '(none)'}")

    host = (parsed.hostname or "").lower()
    if not host:
        raise ValueError(f"URL has no host: {url!r}")

    try:
        infos = socket.getaddrinfo(host, None)
    except OSError as exc:
        raise ValueError(f"Could not resolve host {host!r}: {exc}") from exc

    addresses = {ipaddress.ip_address(info[4][0]) for info in infos}
    if not addresses:
        raise ValueError(f"Host {host!r} resolved to no addresses")

    for addr in addresses:
        if not addr.is_global:
            raise ValueError(
                f"Host {host!r} resolves to non-public address {addr}. "
                f"Only public web pages are supported."
            )

    return host, str(next(iter(addresses)))


class _PinnedDNSAdapter(HTTPAdapter):
    """HTTPAdapter that connects to a pinned IP address.

    DNS pinning prevents the TOCTOU race where ``_resolve_public_host``
    validates a public address and the connection pool independently
    re-resolves the hostname moments later — an attacker with a short-TTL
    DNS record can swap the address after validation.

    The adapter sets ``server_hostname`` and ``assert_hostname`` so that
    TLS SNI and certificate verification still use the original hostname
    even though the TCP connection targets a raw IP.
    """

    def __init__(self, server_hostname: str, *args, **kwargs):
        self._assert_host = server_hostname
        super().__init__(*args, **kwargs)

    def init_poolmanager(self, *args, **kwargs):
        kwargs["server_hostname"] = self._assert_host
        kwargs["assert_hostname"] = self._assert_host
        super().init_poolmanager(*args, **kwargs)


def _pinned_url(url: str, hostname: str, ip_address: str) -> str:
    """Return *url* with its host component replaced by *ip_address*.

    Rebuilds the netloc from parsed components rather than doing a
    substring replace: ``urlparse().hostname`` is always lowercased, so a
    text ``url.replace(hostname, ...)`` silently no-ops on any URL whose
    host isn't already all-lowercase (e.g. ``https://Example.com/``) —
    the original hostname would then reach ``requests``, which re-resolves
    DNS at connect time and reopens the exact TOCTOU gap this pin exists
    to close. Preserves userinfo and port; brackets IPv6 literals.
    """
    parsed = urlparse(url)
    userinfo = ""
    if "@" in parsed.netloc:
        userinfo = parsed.netloc.rsplit("@", 1)[0] + "@"
    host_literal = f"[{ip_address}]" if ":" in ip_address else ip_address
    port = f":{parsed.port}" if parsed.port else ""
    netloc = f"{userinfo}{host_literal}{port}"
    return parsed._replace(netloc=netloc).geturl()


def fetch_and_extract(url: str) -> str:
    """Fetch *url* and return the readable text content.

    Uses readability-lxml to strip navigation, ads, and boilerplate.
    Raises ``ValueError`` for bad URLs, non-public hosts, or network
    failures.

    Redirects are followed manually — each hop is validated through
    :func:`_resolve_public_host` and the TCP connection is pinned to the
    specific IP returned by that validation, so a public host that
    redirects to an internal address is refused and a DNS-rebinding
    attack cannot swap the address between validation and connect time.
    """
    current_url: str = url

    for hop in range(_MAX_REDIRECTS + 1):
        hostname, pinned_ip = _resolve_public_host(current_url)

        adapter = _PinnedDNSAdapter(hostname)
        session = requests.Session()
        session.headers.update({"User-Agent": _USER_AGENT})
        session.mount("https://", adapter)
        session.mount("http://", adapter)

        try:
            resp = session.get(
                _pinned_url(current_url, hostname, pinned_ip),
                timeout=_REQUEST_TIMEOUT,
                allow_redirects=False,
            )
        except requests.ConnectionError as exc:
            raise ValueError(f"Could not connect to {current_url}: {exc}") from exc
        except requests.Timeout as exc:
            raise ValueError(f"Request timed out after {_REQUEST_TIMEOUT}s: {current_url}") from exc
        except requests.RequestException as exc:
            raise ValueError(f"Failed to fetch {current_url}: {exc}") from exc

        if resp.is_redirect:
            if hop >= _MAX_REDIRECTS:
                raise ValueError(f"Too many redirects (max {_MAX_REDIRECTS}) from {url}")
            location = resp.headers.get("Location")
            if not location:
                raise ValueError(f"Redirect from {current_url} with no Location header")
            current_url = urljoin(current_url, location)
            continue

        try:
            resp.raise_for_status()
        except requests.HTTPError as exc:
            raise ValueError(f"HTTP error fetching {current_url}: {exc}") from exc

        break
    else:
        raise ValueError(f"Too many redirects (max {_MAX_REDIRECTS}) from {url}")

    content_type = resp.headers.get("Content-Type", "")
    if "pdf" in content_type.lower():
        raise ValueError("The URL points to a PDF file. Please provide a link to an HTML page.")

    raw = resp.content[:_MAX_RESPONSE_BYTES]
    if len(resp.content) > _MAX_RESPONSE_BYTES:
        logger.warning(
            "Response from %s exceeded %d bytes; truncated",
            current_url,
            _MAX_RESPONSE_BYTES,
        )

    try:
        from readability import Document as ReadabilityDocument
    except ImportError:
        raise ImportError(
            "readability-lxml is required for URL import. "
            "Install it with: pip install readability-lxml"
        )

    try:
        doc = ReadabilityDocument(raw.decode("utf-8", errors="replace"))
        html_content = doc.summary()
    except Exception as exc:
        raise ValueError(f"Could not extract readable content from {current_url}: {exc}") from exc

    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError(
            "beautifulsoup4 is required for URL import. Install it with: pip install beautifulsoup4"
        )

    soup = BeautifulSoup(html_content, "html.parser")
    text = soup.get_text(separator="\n", strip=True)

    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    if len(text) < 50:
        raise ValueError(
            "The page contained very little readable text. "
            "Please check the URL points to a style guide or author guidelines page."
        )

    return text


def parse_guide_with_llm(extracted_text: str, config: AppConfig | None = None) -> StyleGuide:
    """Send extracted text to the LLM and parse the response into a StyleGuide.

    Routes through :func:`model_harness.driver.run_structured` so the
    response is schema-validated and the endpoint is checked against
    ``EndpointPolicy``.

    Raises ``ValueError`` if the LLM returns unparseable output.
    """
    if config is None:
        config = AppConfig()

    user_prompt = (
        "Below is the extracted text from a journal's author guidelines page. "
        "Parse it into the StyleGuide JSON schema.\n\n"
        f"--- BEGIN EXTRACTED TEXT ---\n{extracted_text}\n--- END EXTRACTED TEXT ---"
    )

    return _parse_guide_via_harness(
        system_prompt=_EXTRACTION_SYSTEM_PROMPT,
        user_prompt=user_prompt,
        config=config,
    )


def preview_guide_from_url(url: str, config: AppConfig | None = None) -> StyleGuide:
    """Fetch a URL, extract content, and parse it into a StyleGuide.

    This is the main entry point: it orchestrates fetch → extract → LLM parse.
    The result is NOT saved; the caller is responsible for showing a review and
    calling ``save_custom_guide()`` if approved.
    """
    text = fetch_and_extract(url)
    return parse_guide_with_llm(text, config)


# ---------------------------------------------------------------------------
# alternative import methods (text, docx, pdf)
# ---------------------------------------------------------------------------


def extract_text_from_docx(path: str) -> str:
    """Extract plain text paragraphs from a .docx file.

    Raises ``ValueError`` if the file cannot be read or contains no text.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to import from .docx files. "
            "Install it with: pip install python-docx"
        )

    try:
        doc = Document(path)
    except Exception as exc:
        raise ValueError(f"Could not read .docx file {path}: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs).strip()
    if len(text) < 50:
        raise ValueError(
            "The .docx file contains very little readable text. "
            "Please check the file contains a style guide or author guidelines."
        )
    return text


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF file.

    Uses PyMuPDF (fitz) — install with ``pip install PyMuPDF``.
    Raises ``ValueError`` if the file cannot be read or contains no text.
    """
    try:
        import fitz
    except ImportError:
        raise ImportError(
            "PyMuPDF is required to import from PDF files. Install it with: pip install PyMuPDF"
        )

    try:
        doc = fitz.open(path)
    except Exception as exc:
        raise ValueError(f"Could not read PDF file {path}: {exc}") from exc

    pages = []
    for page_num in range(len(doc)):
        text = doc[page_num].get_text().strip()
        if text:
            pages.append(text)
    doc.close()

    text = "\n\n".join(pages).strip()
    if len(text) < 50:
        raise ValueError(
            "The PDF file contains very little readable text. "
            "Please check the file contains a style guide or author guidelines."
        )
    return text


def preview_guide_from_text(text: str, config: AppConfig | None = None) -> StyleGuide:
    """Parse raw text directly into a StyleGuide via LLM.

    The result is NOT saved; the caller is responsible for showing a review
    and calling ``save_custom_guide()`` if approved.
    """
    text = text.strip()
    if len(text) < 50:
        raise ValueError(
            "The provided text is too short. "
            "Please paste at least 50 characters of a style guide or author guidelines."
        )
    return parse_guide_with_llm(text, config)


def preview_guide_from_file(path: str, config: AppConfig | None = None) -> StyleGuide:
    """Read a .docx or .pdf file, extract text, and parse into a StyleGuide.

    The result is NOT saved; the caller is responsible for showing a review
    and calling ``save_custom_guide()`` if approved.
    """
    lower = path.lower()
    if lower.endswith(".docx"):
        text = extract_text_from_docx(path)
    elif lower.endswith(".pdf"):
        text = extract_text_from_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {path}. Please provide a .docx or .pdf file.")
    return parse_guide_with_llm(text, config)


# ---------------------------------------------------------------------------
# internal: harness-based LLM call
# ---------------------------------------------------------------------------


def _parse_guide_via_harness(system_prompt: str, user_prompt: str, config: AppConfig) -> StyleGuide:
    """Send a structured extraction request through the model harness.

    Replaces the old :func:`_send_llm_request` / :func:`_send_ollama` /
    :func:`_send_openai` / :func:`_send_anthropic` quartet, which each
    issued a raw HTTP request without endpoint validation.  The harness
    validates the endpoint through ``EndpointPolicy`` before any network
    call, and the response is schema-validated against
    :class:`_GuideExtractionShape`.
    """
    policy = ConcreteEndpointPolicy()
    adapter = _build_harness_adapter(config, policy=policy)
    schema_json = _GuideExtractionShape.model_json_schema()

    provider_str = _provider_string(config)
    endpoint = _harness_endpoint(config)
    api_key = _harness_api_key(config)
    model = config.active_model

    model_config = ModelConnectorConfig(
        provider=provider_str,  # type: ignore[arg-type]
        endpoint=endpoint,
        model=model,
        api_key=api_key,
        timeout_s=120.0,
    )

    request = StructuredRequest(
        instructions=system_prompt,
        input=user_prompt,
        schema_json=schema_json,
        mode=StructuredOutputMode.PROMPTED,
        config=model_config,
    )

    async def _run() -> StyleGuide:
        try:
            result: HarnessResult[_GuideExtractionShape] = await run_structured(
                request,
                adapter,
                _GuideExtractionShape,
                endpoint_policy=policy,
            )
        except (StructuredOutputUnsupported, SchemaValidationFailed) as exc:
            raise ValueError(f"The LLM could not produce a valid style guide: {exc}") from exc
        except Exception as exc:
            raise ValueError(f"Unexpected error during style guide extraction: {exc}") from exc

        logger.info(
            "Style guide extraction: mode_used=%s repaired=%s",
            result.mode_used.value,
            result.repaired,
        )

        # Convert the validated Pydantic model back to a StyleGuide dataclass.
        return StyleGuide.from_dict(result.data.model_dump())

    # ── Sync / async boundary ────────────────────────────────────────────
    try:
        return asyncio.run(_run())
    except RuntimeError as exc:
        if "cannot be called from a running event loop" in str(exc):
            raise RuntimeError(
                "parse_guide_with_llm was invoked from within a running "
                "asyncio event loop.  This function is a synchronous wrapper; "
                "use the async harness path directly instead."
            ) from exc
        raise


# -- Harness adapter / config helpers ------------------------------------------


def _provider_string(config: AppConfig) -> str:
    """Map draft's :class:`LLMProvider` enum to a harness ``Provider`` literal."""
    if config.llm_provider == LLMProvider.ANTHROPIC:
        return "anthropic"
    if config.llm_provider == LLMProvider.OPENAI:
        return "generic-api"
    return "ollama"


def _harness_endpoint(config: AppConfig) -> str:
    """Return the base endpoint URL for the active provider.

    The adapter appends ``/chat/completions`` (OpenAI) or ``/v1/messages``
    (Anthropic), so the endpoint here is the base before those path segments.
    """
    if config.llm_provider == LLMProvider.ANTHROPIC:
        return "https://api.anthropic.com"
    if config.llm_provider == LLMProvider.OPENAI:
        return config.openai_base_url.rstrip("/")
    return config.base_url.rstrip("/")


def _harness_api_key(config: AppConfig) -> str | None:
    """Return the API key for the active provider, or None."""
    if config.llm_provider == LLMProvider.ANTHROPIC:
        return config.anthropic_api_key or None
    if config.llm_provider == LLMProvider.OPENAI:
        return config.openai_api_key or None
    return None


def _build_harness_adapter(
    config: AppConfig,
    policy: ConcreteEndpointPolicy | None = None,
) -> OpenAIProvider | AnthropicProvider:
    """Build the correct harness provider adapter for *config*.

    Both Ollama and OpenAI use an OpenAI-compatible chat-completions
    protocol, so they share :class:`OpenAIProvider`.  Anthropic uses
    :class:`AnthropicProvider`.
    """
    if config.llm_provider == LLMProvider.ANTHROPIC:
        return AnthropicProvider(
            endpoint_policy=policy,
            max_tokens=config.num_ctx,
        )
    return OpenAIProvider(
        provider_type="ollama",  # both Ollama and OpenAI support json_object
        endpoint_policy=policy,
    )
