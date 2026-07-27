"""Shared utilities for writing plain .docx files from paragraph data."""

from __future__ import annotations

import io
import logging
import os
import zipfile

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}

STYLE_MAP: dict[str, str] = {
    "Heading 1": "Heading 1",
    "Heading 2": "Heading 2",
    "Heading 3": "Heading 3",
    "Normal": "Normal",
}


def _add_inline_image(run, blob: bytes, filename: str, content_type: str) -> None:
    """Embed an image blob as an inline drawing inside a run using a temp file."""
    ext = os.path.splitext(filename)[1] or ".png"
    tmp_dir = "C:/Users/mjcas/AppData/Local/Temp/opencode"
    os.makedirs(tmp_dir, exist_ok=True)
    tmp_path = os.path.join(tmp_dir, f"_img_{hash(blob) % 10**8}{ext}")
    with open(tmp_path, "wb") as f:
        f.write(blob)
    try:
        run.add_picture(tmp_path, width=Pt(96))
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def _paragraph_has_images(entry: dict) -> bool:
    images = entry.get("images") or []
    return any(img.get("blob") for img in images)


def write_plain_docx(paragraphs: list[dict], path: str) -> None:
    """Write a plain .docx without track changes (fallback for no-change case).

    Preserves style, indentation, bold, italic, underline, font metadata,
    and inline images.
    """
    logger.debug("Writing plain .docx to %s (%d paragraphs)", path, len(paragraphs))
    doc = Document()

    for entry in paragraphs:
        p = doc.add_paragraph()

        style_name = entry.get("style_name", "Normal")
        pn = STYLE_MAP.get(style_name, "Normal")
        if hasattr(p, "style"):
            try:
                p.style = doc.styles[pn]
            except KeyError:
                pass

        indent = entry.get("indent_level", 0) or 0
        if indent > 0:
            p.paragraph_format.left_indent = Pt(indent)

        if entry.get("space_before") is not None:
            p.paragraph_format.space_before = Pt(entry["space_before"])
        if entry.get("space_after") is not None:
            p.paragraph_format.space_after = Pt(entry["space_after"])
        if entry.get("line_spacing") is not None:
            p.paragraph_format.line_spacing = entry["line_spacing"]

        images = entry.get("images") or []
        text = entry.get("text", "")
        has_images = any(img.get("blob") for img in images)

        if has_images and text:
            run = p.add_run(text)
            if entry.get("is_bold"):
                run.bold = True
            if entry.get("is_italic"):
                run.italic = True
            if entry.get("is_underline"):
                run.underline = True
            if entry.get("font_size"):
                run.font.size = Pt(entry["font_size"])
            if entry.get("font_name"):
                run.font.name = entry["font_name"]
            for img in images:
                if img.get("blob"):
                    _add_inline_image(run, img["blob"], img.get("filename", "image.png"), img.get("content_type", ""))
        elif has_images:
            run = p.add_run("")
            for img in images:
                if img.get("blob"):
                    _add_inline_image(run, img["blob"], img.get("filename", "image.png"), img.get("content_type", ""))
        else:
            run = p.add_run(text)
            if entry.get("is_bold"):
                run.bold = True
            if entry.get("is_italic"):
                run.italic = True
            if entry.get("is_underline"):
                run.underline = True
            if entry.get("font_size"):
                run.font.size = Pt(entry["font_size"])
            if entry.get("font_name"):
                run.font.name = entry["font_name"]

    doc.save(path)
