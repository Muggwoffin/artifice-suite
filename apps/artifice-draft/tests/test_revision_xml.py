"""Regression tests for tracked-changes revision XML.

Verifies that:
  - edited paragraphs produce <w:ins>/<w:del> in the output .docx
  - unchanged paragraphs are preserved exactly
"""

from __future__ import annotations

import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document

from src.doc_parser import parse_docx
from src.doc_writer import apply_edits_to_docx


def _read_doc_xml(path):
    """Extract document.xml text from a .docx ZIP archive."""
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def test_revision_xml_has_ins_and_del(tmp_path: Path):
    """Output docx must contain w:ins or w:del elements when edits exist."""
    doc = Document()
    doc.add_paragraph("This has a typo in it")
    doc.add_paragraph("Second paragraph is correct")
    inp = tmp_path / "rev_input.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: "This is corrected now", 1: None}

    out = tmp_path / "rev_output.docx"
    apply_edits_to_docx(None, paragraphs, edits, str(out))
    assert out.exists()

    doc_xml = _read_doc_xml(str(out))
    root = ET.fromstring(doc_xml)
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    ins_elements = root.findall(".//w:ins", nsmap)
    del_elements = root.findall(".//w:del", nsmap)

    assert len(ins_elements) >= 1, f"Expected at least one <w:ins>, found {len(ins_elements)}"
    assert len(del_elements) >= 1, f"Expected at least one <w:del>, found {len(del_elements)}"


def test_revision_xml_unchanged_paragraph_preserved(tmp_path: Path):
    """The unchanged paragraph must remain exactly as in the original."""
    doc = Document()
    doc.add_paragraph("This has a typo in it")
    doc.add_paragraph("Second paragraph is correct")
    inp = tmp_path / "rev_preserve.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: "Changed text", 1: None}

    out = tmp_path / "rev_preserved_out.docx"
    apply_edits_to_docx(None, paragraphs, edits, str(out))
    assert out.exists()

    doc_xml = _read_doc_xml(str(out))
    root = ET.fromstring(doc_xml)
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    preserved_texts = []
    for p in root.findall(".//w:p", nsmap):
        for r in p.findall("w:r", nsmap):
            t_el = r.find("w:t", nsmap)
            if t_el is not None:
                preserved_texts.append(t_el.text or "")

    assert "Second paragraph is correct" in preserved_texts, \
        f"Expected unchanged text 'Second paragraph is correct' not found. Got: {preserved_texts}"


def test_revision_xml_with_real_input_path(tmp_path: Path):
    """Passing a real input_path should preserve original formatting."""
    doc = Document()
    doc.add_paragraph("Original text here")
    doc.add_paragraph("Keep this unchanged")
    inp = tmp_path / "rev_real_input.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: "Corrected text here"}

    out = tmp_path / "rev_with_input.docx"
    apply_edits_to_docx(str(inp), paragraphs, edits, str(out))
    assert out.exists()

    doc_xml = _read_doc_xml(str(out))
    assert "Corrected text here" in doc_xml
    assert "Keep this unchanged" in doc_xml
