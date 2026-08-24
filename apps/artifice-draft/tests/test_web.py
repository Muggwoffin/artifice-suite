# SPDX-FileCopyrightText: 2026 Maurice Casey
#
# SPDX-License-Identifier: AGPL-3.0-or-later

"""Tests for the FastAPI web frontend (src/web/).

Mocks ``asyncio.run`` inside ``call_ollama`` so no real Ollama/OpenAI/Anthropic
request ever leaves the test process.
"""

from __future__ import annotations

import json
import time
from unittest.mock import patch

import pytest
from docx import Document
from fastapi import HTTPException
from fastapi.testclient import TestClient

import artifice_draft.web.runtime as runtime
from artifice_draft.llm_edit import LLMEdit
from artifice_draft.web.server import app


@pytest.fixture(autouse=True)
def isolated_settings(tmp_path, monkeypatch):
    """Never let a test touch the real ~/.artifice_draft/web_settings.json.

    OCR Pipeline (the sibling project this web build is modeled on) shipped a
    test that skipped this once and wrote real content into a developer's
    actual settings file; the fix there was making this the standard fixture
    pattern for every test that touches persisted settings, and this project
    follows the same rule from the start.
    """
    monkeypatch.setattr(runtime, "_SETTINGS_PATH", tmp_path / "web_settings.json")


@pytest.fixture()
def client():
    return TestClient(app)


@pytest.fixture()
def docx_bytes(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello wrold")
    doc.add_paragraph("This is a secnd paragraph.")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return p.read_bytes()


def _wait_for_thread(doc_id: str, timeout: float = 5.0) -> None:
    doc = runtime.state.get(doc_id)
    assert doc is not None
    assert doc.thread is not None
    doc.thread.join(timeout=timeout)


def _mock_edits_response():
    """Return a list of LLMEdit objects matching the mocked harness output."""
    return [
        LLMEdit(paragraph_index=0, original_text="Hello wrold", edited_text="Hello world", status="edited"),
        LLMEdit(paragraph_index=1, original_text="This is a secnd paragraph.", edited_text=None, status="unchanged"),
    ]


# --------------------------------------------------------------------- upload

def test_upload_rejects_non_docx(client):
    res = client.post("/api/upload", files={"file": ("notes.txt", b"hello", "text/plain")})
    assert res.status_code == 400


def test_upload_parses_paragraphs(client, docx_bytes):
    res = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    assert res.status_code == 200
    body = res.json()
    assert body["paragraph_count"] == 2
    assert body["stage"] == "uploaded"
    assert body["doc_id"]


# ------------------------------------------------------------------- settings

def test_settings_roundtrip_and_merge(client):
    res = client.get("/api/settings")
    assert res.status_code == 200
    defaults = res.json()
    assert defaults["llm_provider"] == "ollama"

    res = client.post("/api/settings", json={"batch_size": 3, "enable_review": True})
    assert res.status_code == 200
    updated = res.json()
    assert updated["batch_size"] == 3
    assert updated["enable_review"] is True
    # author_name untouched by the partial patch — merge, not replace.
    assert updated["author_name"] == defaults["author_name"]

    res = client.post("/api/settings", json={"author_name": "Test Author"})
    updated2 = res.json()
    assert updated2["author_name"] == "Test Author"
    # earlier fields survive the second, unrelated patch.
    assert updated2["batch_size"] == 3
    assert updated2["enable_review"] is True


def test_settings_never_expose_api_keys(client):
    res = client.get("/api/settings")
    body = res.json()
    assert "openai_api_key" not in body
    assert "anthropic_api_key" not in body


# ------------------------------------------------------------- run: no review

def test_run_without_review_writes_output_and_is_downloadable(client, docx_bytes):
    client.post("/api/settings", json={"enable_review": False})

    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]

    # resolve_for_run probes a live Ollama to pick a model; these tests
    # exercise the run/download flow, not resolution, so it is mocked
    # exactly as the LLM call below already is.
    with (
        patch("artifice_draft.web.runtime.resolve_for_run"),
        patch("artifice_draft.llm_client.asyncio.run", return_value=_mock_edits_response()),
    ):
        start = client.post(f"/api/run/{doc_id}/start")
        assert start.status_code == 200
        _wait_for_thread(doc_id)

    status = client.get(f"/api/run/{doc_id}/status").json()
    assert status["stage"] == "done"
    assert status["output_filename"]

    download = client.get(f"/api/run/{doc_id}/download")
    assert download.status_code == 200
    assert len(download.content) > 0


def test_download_404s_before_finalize(client, docx_bytes):
    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]
    res = client.get(f"/api/run/{doc_id}/download")
    assert res.status_code == 404


def test_unknown_doc_id_404s(client):
    assert client.get("/api/run/doesnotexist/status").status_code == 404
    assert client.post("/api/run/doesnotexist/start").status_code == 404
    assert client.get("/api/run/doesnotexist/review").status_code == 404
    assert client.get("/api/run/doesnotexist/download").status_code == 404


# ---------------------------------------------------------------- run: review

def test_review_flow_reject_keeps_original_text(client, docx_bytes):
    client.post("/api/settings", json={"enable_review": True})

    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]

    # resolve_for_run probes a live Ollama to pick a model; these tests
    # exercise the run/download flow, not resolution, so it is mocked
    # exactly as the LLM call below already is.
    with (
        patch("artifice_draft.web.runtime.resolve_for_run"),
        patch("artifice_draft.llm_client.asyncio.run", return_value=_mock_edits_response()),
    ):
        client.post(f"/api/run/{doc_id}/start")
        _wait_for_thread(doc_id)

    status = client.get(f"/api/run/{doc_id}/status").json()
    assert status["stage"] == "awaiting_review"

    review = client.get(f"/api/run/{doc_id}/review").json()
    # Only the actually-changed paragraph appears — the "unchanged" one from
    # the mocked response is filtered out, same rule cli_review() applies.
    assert len(review["items"]) == 1
    item = review["items"][0]
    assert item["paragraph_index"] == 0
    assert item["original_text"] == "Hello wrold"
    assert item["edited_text"] == "Hello world"
    assert item["diff"]["original_ranges"]

    submit = client.post(f"/api/run/{doc_id}/review", json={
        "decisions": [{"paragraph_index": 0, "approved": False, "replacement_text": None}],
    })
    assert submit.status_code == 200
    assert submit.json()["stage"] == "done"

    download = client.get(f"/api/run/{doc_id}/download")
    assert download.status_code == 200


def test_review_endpoint_409s_when_not_awaiting_review(client, docx_bytes):
    upload = client.post("/api/upload", files={
        "file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    doc_id = upload.json()["doc_id"]
    res = client.get(f"/api/run/{doc_id}/review")
    assert res.status_code == 409


# --------------------------------------------------------- credential redaction

def test_get_settings_redacts_api_key(client, monkeypatch, tmp_path):
    """GET /api/settings must return the placeholder, not the real key."""
    from artifice_draft.web import runtime as rt_module

    monkeypatch.setattr(rt_module, "_SETTINGS_PATH", tmp_path / "web_settings.json")
    rt_module.save_settings({"api_key": "sk-real-secret", "model_name": "test-model"})

    res = client.get("/api/settings")
    assert res.status_code == 200
    body = res.json()
    assert body["api_key"] == "*" * 12
    assert body["model_name"] == "test-model"


def test_settings_roundtrip_preserves_api_key(client, monkeypatch, tmp_path):
    """GET then POST the response back unmodified — the stored key must
    survive, not be overwritten with the redacted placeholder."""
    from artifice_draft.web import runtime as rt_module

    monkeypatch.setattr(rt_module, "_SETTINGS_PATH", tmp_path / "web_settings.json")
    rt_module.save_settings({"api_key": "sk-original-key-abc"})

    # Step 1: GET — receives the redacted placeholder.
    res_get = client.get("/api/settings")
    assert res_get.status_code == 200
    body = res_get.json()
    assert body["api_key"] == "*" * 12  # redacted on the way out

    # Step 2: POST the exact same body back.
    res_post = client.post("/api/settings", json=body)
    assert res_post.status_code == 200

    # Step 3: Verify the stored key is still the original.
    stored = rt_module.load_settings()
    assert stored["api_key"] == "sk-original-key-abc", (
        "The stored key must survive a round-trip of the redacted placeholder"
    )


# --------------------------------------------------------------- upload limits

def test_upload_rejects_oversized_content_length(client):
    """A Content-Length header exceeding 50 MB must return HTTP 413."""
    res = client.post(
        "/api/upload",
        files={"file": ("large.docx", b"x" * 100, "application/octet-stream")},
        headers={"Content-Length": str(51 * 1024 * 1024)},
    )
    assert res.status_code == 413


def test_upload_accepts_exactly_at_limit(client, docx_bytes):
    """A file at exactly 50 MB (Content-Length) must be accepted."""
    res = client.post(
        "/api/upload",
        files={"file": ("sample.docx", docx_bytes,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        headers={"Content-Length": str(50 * 1024 * 1024)},
    )
    # Should pass the size check — may fail on content if the file isn't
    # actually 50 MB, but the 413 must not fire.
    assert res.status_code != 413


def test_style_guide_preview_rejects_oversized(client):
    """preview-file route also rejects oversized uploads."""
    res = client.post(
        "/api/style-guides/preview-file",
        files={"file": ("large.pdf", b"x" * 100, "application/pdf")},
        headers={"Content-Length": str(51 * 1024 * 1024)},
    )
    assert res.status_code == 413


# ------------------------------------------------------ streaming cap (bypass)

def test_read_capped_raises_during_read():
    """_read_capped raises HTTP 413 once the limit is exceeded mid-stream,
    before the full body is gathered."""
    import asyncio
    from artifice_draft.web.server import _read_capped

    class _FakeUpload:
        filename = "test.docx"
        def __init__(self, total: int):
            self._remain = total
        async def read(self, size: int = -1) -> bytes:
            if self._remain <= 0:
                return b""
            n = min(size if size > 0 else 4096, 4096, self._remain)
            self._remain -= n
            return b"x" * n

    with pytest.raises(HTTPException) as exc_info:
        asyncio.run(_read_capped(_FakeUpload(10_000), 10))
    assert exc_info.value.status_code == 413


def test_upload_streaming_cap_catches_what_content_length_misses(client, monkeypatch):
    """When the Content-Length check passes (header absent), the streaming
    cap in _read_capped is still enforced and the upload is refused.

    The existing ``test_upload_rejects_oversized_content_length`` proves the
    Content-Length fast path works for honest clients.  This test proves the
    streaming cap works when the fast path is skipped — exactly the hole a
    chunked-transfer upload exploits.
    """
    import artifice_draft.web.server as server_mod

    # Force Content-Length check to always pass (simulating absent header).
    monkeypatch.setattr(server_mod, "_content_length_exceeds", lambda *a, **kw: False)
    monkeypatch.setattr(server_mod, "_MAX_UPLOAD_BYTES", 10)

    res = client.post("/api/upload", files={
        "file": ("test.docx", b"x" * 1000,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    })
    assert res.status_code == 413


def test_preview_file_streaming_cap(client, monkeypatch):
    """Same bypass guard on the /api/style-guides/preview-file route."""
    import artifice_draft.web.server as server_mod

    monkeypatch.setattr(server_mod, "_content_length_exceeds", lambda *a, **kw: False)
    monkeypatch.setattr(server_mod, "_MAX_UPLOAD_BYTES", 10)

    res = client.post("/api/style-guides/preview-file", files={
        "file": ("test.pdf", b"x" * 1000, "application/pdf"),
    })
    assert res.status_code == 413
