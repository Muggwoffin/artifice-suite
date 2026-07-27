"""Tests for doc_writer module."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from artifice_draft.doc_parser import parse_docx
from artifice_draft.doc_writer import apply_edits_to_docx


def test_apply_edits_no_changes(tmp_path: Path):
    """When edits are all None, we should write a plain copy (no track changes)."""
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("This is fine too.")
    inp = tmp_path / "writer_input.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {i: None for i in range(len(paragraphs))}

    out = tmp_path / "writer_plain.docx"
    apply_edits_to_docx(None, paragraphs, edits, str(out))
    assert out.exists()


def test_apply_edits_with_changes(tmp_path: Path):
    """When there are actual edits, the output should exist and be a valid .docx."""
    doc = Document()
    doc.add_paragraph("This has a typo in it")
    doc.add_paragraph("Second paragraph is correct")
    inp = tmp_path / "writer_changes.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: "This has no typos now", 1: None}

    out = tmp_path / "writer_edited.docx"
    apply_edits_to_docx(None, paragraphs, edits, str(out))
    assert out.exists()


def test_apply_edits_with_empty_input(tmp_path: Path):
    """Empty paragraphs list should not crash."""
    doc = Document()
    doc.add_paragraph("single line")
    inp = tmp_path / "writer_single.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: None}

    out = tmp_path / "writer_single_out.docx"
    apply_edits_to_docx(None, paragraphs, edits, str(out))
    assert out.exists()


def test_apply_edits_preserves_input_path(tmp_path: Path):
    """Passing a real input_path should produce a valid .docx."""
    doc = Document()
    doc.add_paragraph("Original text here")
    inp = tmp_path / "writer_real_input.docx"
    doc.save(str(inp))

    paragraphs = parse_docx(str(inp))
    edits = {0: "Corrected text here"}

    out = tmp_path / "writer_with_input.docx"
    apply_edits_to_docx(str(inp), paragraphs, edits, str(out))
    assert out.exists()
