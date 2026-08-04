# SPDX-FileCopyrightText: 2026 Maurice Casey
#
# SPDX-License-Identifier: AGPL-3.0-or-later

"""Shared test fixtures for the copy-edit test suite."""

from __future__ import annotations

import pytest
from docx import Document


@pytest.fixture()
def sample_paragraphs() -> list[dict]:
    """Two standard test paragraphs with all new fields."""
    return [
        {
            "paragraph_index": 0, "text": "Hello world",
            "style_name": "Normal",
            "is_bold": False, "is_italic": False, "is_underline": False,
            "indent_level": 0, "font_size": None, "font_name": None,
            "alignment": None, "space_before": None, "space_after": None,
            "line_spacing": None, "is_list_item": False, "list_level": 0,
            "language": None,
        },
        {
            "paragraph_index": 1, "text": "This is a second paragraph.",
            "style_name": "Normal",
            "is_bold": False, "is_italic": False, "is_underline": False,
            "indent_level": 0, "font_size": None, "font_name": None,
            "alignment": None, "space_before": None, "space_after": None,
            "line_spacing": None, "is_list_item": False, "list_level": 0,
            "language": None,
        },
    ]


@pytest.fixture()
def sample_docx(tmp_path):
    """Create a minimal .docx with two paragraphs and return its path."""
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("This is a second paragraph.")
    p = tmp_path / "sample.docx"
    doc.save(str(p))
    return str(p)
