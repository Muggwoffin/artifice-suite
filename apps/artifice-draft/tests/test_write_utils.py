"""Tests for write_utils: tempfile usage instead of hardcoded paths — finding #4."""

from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from artifice_draft.write_utils import _add_inline_image


class TestAddInlineImageTempfile:
    """Finding #4: _add_inline_image must use tempfile, not a hardcoded path."""

    def test_uses_system_temp_dir(self):
        """The temp file created by _add_inline_image must live under the
        system temp directory, not a hardcoded user-specific path."""
        doc = Document()
        run = doc.add_paragraph().add_run("x")
        blob = b"\x89PNG\r\n\x1a\n" + b"\x00" * 100  # minimal fake PNG header

        import tempfile as tf

        # Capture the path that NamedTemporaryFile creates.
        captured_paths = []

        orig_named_temp = tf.NamedTemporaryFile

        def _capturing_temp(*args, **kwargs):
            f = orig_named_temp(*args, **kwargs)
            captured_paths.append(f.name)
            return f

        # Mock add_picture so it doesn't try to parse the fake PNG blob.
        with (
            patch("artifice_draft.write_utils.tempfile.NamedTemporaryFile", _capturing_temp),
            patch.object(run, "add_picture", MagicMock()),
        ):
            _add_inline_image(run, blob, "test.png", "image/png")

        assert len(captured_paths) == 1
        temp_path = captured_paths[0]
        # The temp file should be under the system temp dir.
        system_tmp = Path(tempfile.gettempdir())
        assert system_tmp in Path(temp_path).parents or temp_path.startswith(str(system_tmp))

        # The file should be cleaned up (unlinked) after add_picture returns.
        assert not Path(temp_path).exists()

    def test_does_not_hardcode_windows_path(self):
        """The hardcoded path must not appear anywhere in the module."""
        import artifice_draft.write_utils as wu
        import inspect
        source = inspect.getsource(wu._add_inline_image)
        assert "C:/Users/mjcas" not in source
        assert "AppData" not in source
        assert "tempfile" in source or "NamedTemporaryFile" in source

    def test_handles_no_extension(self):
        """When filename has no extension, defaults to .png."""
        doc = Document()
        run = doc.add_paragraph().add_run("x")
        blob = b"\x89PNG\r\n\x1a\n" + b"\x00" * 100

        import tempfile as tf

        captured_paths = []

        orig_named_temp = tf.NamedTemporaryFile

        def _capturing_temp(*args, **kwargs):
            f = orig_named_temp(*args, **kwargs)
            captured_paths.append(f.name)
            return f

        with (
            patch("artifice_draft.write_utils.tempfile.NamedTemporaryFile", _capturing_temp),
            patch.object(run, "add_picture", MagicMock()),
        ):
            _add_inline_image(run, blob, "noextension", "image/png")

        assert len(captured_paths) == 1
        assert captured_paths[0].endswith(".png")
